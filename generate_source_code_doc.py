import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_code_document():
    doc = docx.Document()

    # ── Page Margins ──
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # ── Styles & Colors ──
    NAVY = RGBColor(13, 17, 23)
    TEAL = RGBColor(34, 150, 140)
    DARK_BLUE = RGBColor(26, 54, 93)
    MUTED = RGBColor(100, 110, 125)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title.add_run("MindSense AI — Complete Source Code Documentation\n")
    run_t.font.name = "Arial"
    run_t.font.size = Pt(22)
    run_t.font.bold = True
    run_t.font.color.rgb = DARK_BLUE

    sub = title.add_run("Multimodal Psychiatric Evaluation & Real-Time Gated Fusion System\nHack4Health Hackathon Project")
    sub.font.name = "Arial"
    sub.font.size = Pt(11)
    sub.font.italic = True
    sub.font.color.rgb = MUTED

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Table of Categories Overview
    categories = [
        ("1. Frontend & Application Layer", [
            ("app.py", "Streamlit 8-Tab Interactive Frontend UI & Live Webcam Controller")
        ]),
        ("2. Core Multimodal Engine & Inference", [
            ("multimodal_pipeline.py", "Core Gated Fusion Engine, PyTorch MLP, Keras CNN & Audio pipeline"),
            ("audio_features.py", "280-Dimensional Acoustic MFCC & Spectral Feature Extractor")
        ]),
        ("3. PyTorch Multitask MLP Package", [
            ("mental_health_model_package/train_best_mlp.py", "PyTorch Multitask Tabular MLP Model Training Script"),
            ("mental_health_model_package/inference.py", "Standalone PyTorch MLP Inference Helper")
        ]),
        ("4. Model Training & Evaluation Pipelines", [
            ("training/phase1_csv_baseline.py", "Phase 1: Tabular Baseline Classification & Regression"),
            ("training/phase2_facial_cnn.py", "Phase 2: Baseline Facial Emotion CNN"),
            ("training/phase2b_facial_cnn_transfer.py", "Phase 2b: BatchNorm & Augmentation Facial CNN"),
            ("training/phase2c_improve_sad_cnn.py", "Phase 2c: Deep VGG Facial CNN with Sad Class Boosting"),
            ("training/phase3_audio_classifier.py", "Phase 3: Initial XGBoost Audio Emotion Classifier"),
            ("training/phase3b_audio_improved.py", "Phase 3b: Improved 280-dim Acoustic Classifier"),
            ("training/phase4_fusion.py", "Phase 4: Late Fusion Layer & Multimodal Evaluation"),
            ("training/phase5_explainability.py", "Phase 5: SHAP Interpretability & Feature Driver Analysis"),
            ("training/eval_on_original_val.py", "Head-to-head Validation Set Evaluator")
        ]),
        ("5. Data Cleaning & API Server", [
            ("clean_dataset.py", "Automated CNN Label-Noise Detection & Image Quarantine"),
            ("backend_api.py", "FastAPI REST API Service Endpoint")
        ]),
        ("6. Infrastructure & Deployment", [
            ("Dockerfile", "Production Container Definition"),
            ("requirements.txt", "Project Python Dependencies"),
            ("runtime.txt", "Streamlit Cloud Python 3.11 Specification")
        ])
    ]

    # Add Summary Table
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Category"
    hdr[1].text = "File Name"
    hdr[2].text = "Module Description"

    for cell in hdr:
        shading = parse_xml(r'<w:shd {} w:fill="1A365D"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.name = "Arial"
                r.font.size = Pt(9.5)

    for cat_name, files in categories:
        for fname, fdesc in files:
            row_cells = table.add_row().cells
            row_cells[0].text = cat_name
            row_cells[1].text = fname
            row_cells[2].text = fdesc
            for c in row_cells:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.name = "Arial"
                        r.font.size = Pt(8.5)

    doc.add_page_break()

    # Helper function to add code file block
    def add_code_file(filepath, cat_name, description):
        if not os.path.exists(filepath):
            return

        h = doc.add_heading(level=2)
        r_h = h.add_run(f"📄 {os.path.basename(filepath)}")
        r_h.font.name = "Arial"
        r_h.font.size = Pt(14)
        r_h.font.bold = True
        r_h.font.color.rgb = DARK_BLUE

        p_meta = doc.add_paragraph()
        p_meta.paragraph_format.space_after = Pt(6)
        r_cat = p_meta.add_run(f"Category: {cat_name}  |  Path: {filepath}\n")
        r_cat.font.name = "Arial"
        r_cat.font.size = Pt(8.5)
        r_cat.font.bold = True
        r_cat.font.color.rgb = TEAL

        r_desc = p_meta.add_run(f"Description: {description}")
        r_desc.font.name = "Arial"
        r_desc.font.size = Pt(8.5)
        r_desc.font.italic = True
        r_desc.font.color.rgb = MUTED

        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            code_text = f.read()

        lines = code_text.splitlines()
        p_code = doc.add_paragraph()
        p_code.paragraph_format.space_before = Pt(4)
        p_code.paragraph_format.space_after = Pt(14)
        p_code.paragraph_format.line_spacing = 1.05

        # Code block background container using single cell table
        c_table = doc.add_table(rows=1, cols=1)
        c_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = c_table.cell(0, 0)
        shading = parse_xml(r'<w:shd {} w:fill="F8F9FA"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)
        
        # Border style
        borders = parse_xml(r'''
            <w:tcBorders {} >
                <w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E0"/>
                <w:left w:val="single" w:sz="12" w:space="0" w:color="1A365D"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E0"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="CBD5E0"/>
            </w:tcBorders>
        '''.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(borders)

        cell_p = cell.paragraphs[0]
        cell_p.paragraph_format.space_before = Pt(4)
        cell_p.paragraph_format.space_after = Pt(4)
        cell_p.paragraph_format.line_spacing = 1.05

        # Format code with line numbers
        for idx, line in enumerate(lines, 1):
            r_num = cell_p.add_run(f"{idx:4d}  ")
            r_num.font.name = "Consolas"
            r_num.font.size = Pt(7.5)
            r_num.font.color.rgb = RGBColor(140, 150, 165)

            r_code = cell_p.add_run(f"{line}\n")
            r_code.font.name = "Consolas"
            r_code.font.size = Pt(8.0)
            r_code.font.color.rgb = NAVY

        doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Process all files by category
    for cat_name, files in categories:
        h1 = doc.add_heading(level=1)
        r1 = h1.add_run(cat_name)
        r1.font.name = "Arial"
        r1.font.size = Pt(16)
        r1.font.bold = True
        r1.font.color.rgb = DARK_BLUE
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

        for fname, fdesc in files:
            add_code_file(fname, cat_name, fdesc)

    output_path = "/Users/aryaagarwal/Downloads/MindSense_AI_Source_Code_Documentation.docx"
    doc.save(output_path)
    print(f"Successfully generated Word Document -> {output_path}")

if __name__ == "__main__":
    create_code_document()
